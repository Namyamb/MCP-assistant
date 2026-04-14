import os
import json
import base64
import urllib.parse
from http.server import ThreadingHTTPServer, SimpleHTTPRequestHandler
from pathlib import Path
from app.core.config import WEB_HOST, WEB_PORT, PROJECT_ROOT, DATA_DIR
from app.integrations.gmail.core import get_emails, get_email_by_id
from app.core.orchestrator import run_agent


STATIC_DIR = PROJECT_ROOT / "app" / "ui" / "static"

# Mime types for images that the vision model can process
_IMAGE_EXTS = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
               '.gif': 'image/gif', '.webp': 'image/webp'}

# Text-like files we can read directly
_TEXT_EXTS = {'.txt', '.md', '.csv', '.json', '.py', '.js', '.ts',
              '.html', '.xml', '.yaml', '.yml', '.log', '.ini', '.toml'}

MAX_TEXT_CHARS = 12_000  # cap injected content so we don't overflow context


def _extract_text(fpath: Path) -> str | None:
    """Return plain-text content of a file, or None if unsupported / binary."""
    ext = fpath.suffix.lower()
    try:
        if ext in _TEXT_EXTS:
            text = fpath.read_text(encoding='utf-8', errors='replace')
            return text[:MAX_TEXT_CHARS] + ("\n...[truncated]" if len(text) > MAX_TEXT_CHARS else "")

        if ext == '.docx':
            from docx import Document
            doc = Document(str(fpath))
            # Include paragraph text AND table cell text
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            text = '\n'.join(parts)
            return text[:MAX_TEXT_CHARS] + ("\n...[truncated]" if len(text) > MAX_TEXT_CHARS else "")

        if ext == '.pdf':
            try:
                import PyPDF2
                with open(fpath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    pages = [p.extract_text() for p in reader.pages if p.extract_text()]
                text = '\n\n'.join(pages)
                return text[:MAX_TEXT_CHARS] + ("\n...[truncated]" if len(text) > MAX_TEXT_CHARS else "")
            except ImportError:
                return "[PDF support needs PyPDF2 — run: pip install PyPDF2]"

    except Exception as e:
        return f"[Could not read file: {e}]"

    return None  # unsupported extension


agent_state = {}


class AgentHTTPRequestHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(STATIC_DIR), **kwargs)

    def do_GET(self):
        parsed_path = urllib.parse.urlparse(self.path)
        if parsed_path.path == '/api/health':
            self.send_json({"status": "ok"})
        elif parsed_path.path == '/api/inbox':
            try:
                emails = get_emails(limit=20)
                self.send_json({"emails": emails})
            except PermissionError as e:
                self.send_error(401, str(e))
            except Exception as e:
                self.send_error(500, str(e))
        elif parsed_path.path == '/api/email':
            query = urllib.parse.parse_qs(parsed_path.query)
            msg_id = query.get('message_id', [''])[0]
            if msg_id:
                sel_email = get_email_by_id(msg_id)
                self.send_json({"email": sel_email})
            else:
                self.send_error(400, "Missing message_id")
        else:
            super().do_GET()

    def do_POST(self):
        global agent_state
        parsed_path = urllib.parse.urlparse(self.path)

        if parsed_path.path == '/api/reset_history':
            try:
                length = int(self.headers.get('Content-Length', 0))
                data = json.loads(self.rfile.read(length).decode('utf-8')) if length else {}
                mode = data.get('mode', 'gmail')
                agent_state.pop(f"history_{mode}", None)
                # Also clear any pending attachment state
                agent_state.pop("last_image", None)
                agent_state.pop("last_attachment_path", None)
                self.send_json({"ok": True})
            except Exception:
                self.send_json({"ok": False})

        elif parsed_path.path == '/api/chat':
            try:
                length = int(self.headers.get('Content-Length', 0))
                if length == 0:
                    self.send_error(400, "Empty request body")
                    return
                data = json.loads(self.rfile.read(length).decode('utf-8'))
                msg  = data.get('message', '')
                attachment = data.get('attachment')
                mode = data.get('mode', 'gmail')

                if attachment:
                    ups = DATA_DIR / "uploads"
                    ups.mkdir(exist_ok=True)
                    fname = attachment['name']
                    fpath = ups / fname
                    fpath.write_bytes(base64.b64decode(attachment['data']))

                    # Always store path so Gmail tools can attach files to emails
                    agent_state["last_attachment_path"] = str(fpath)

                    ext = Path(fname).suffix.lower()

                    if ext in _IMAGE_EXTS:
                        # Store for vision processing in orchestrator
                        agent_state["last_image"] = {
                            "data": attachment['data'],   # raw base64 string
                            "mime": _IMAGE_EXTS[ext],
                            "name": fname
                        }
                        msg += f"\n[User attached image: {fname}]"
                    else:
                        # Extract text content and inject into the message
                        content = _extract_text(fpath)
                        if content:
                            msg += f"\n\n--- Attached file: {fname} ---\n{content}\n--- End of file ---"
                        else:
                            msg += f"\n[Attached binary file: {fname} — content cannot be displayed]"

                from app.main import build_server
                mcp   = build_server()
                reply = run_agent(msg, mcp, agent_state, mode=mode)
                self.send_json({"reply": reply})

            except Exception as e:
                import traceback
                traceback.print_exc()
                self.send_json({"reply": f"<b>System Error:</b> {str(e)}"})
        else:
            self.send_error(404, "Not Found")

    def send_json(self, data):
        self.send_response(200)
        self.send_header('Content-type', 'application/json')
        self.end_headers()
        self.wfile.write(json.dumps(data).encode('utf-8'))

    def log_message(self, fmt, *args):
        pass  # suppress request logs in console


def run_web():
    server = ThreadingHTTPServer((WEB_HOST, WEB_PORT), AgentHTTPRequestHandler)
    print(f"Starting web server on http://{WEB_HOST}:{WEB_PORT}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nShutting down server.")
        server.server_close()
